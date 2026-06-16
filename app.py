import streamlit as st
import pandas as pd
import json
import os
import io
import copy
import subprocess
import platform
import zipfile
from docx import Document
from docx.text.paragraph import Paragraph

# ==========================================
# 1. CONFIGURAÇÃO DA PÁGINA E CRÉDITOS
# ==========================================
st.set_page_config(
    page_title="Etiquetas Lichens",
    page_icon="icone.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
    .titulo { color: #10b981; font-weight: bold; }
    .autor { color: #64748b; font-size: 0.9em; font-style: italic; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<h1 class="titulo">🔬 Etiquetas Lichens</h1>', unsafe_allow_html=True)
st.markdown('<p class="autor">Desenvolvido por: Kristian Lancelot | Laboratório de Liquenologia</p>', unsafe_allow_html=True)
st.divider()

# ==========================================
# 2. FUNÇÕES AUXILIARES
# ==========================================
CONFIG_FILE = "preferencias_colunas.json"

def carregar_preferencias():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r') as f: return json.load(f)
    return {}

def guardar_preferencias(preferencias):
    with open(CONFIG_FILE, 'w') as f: json.dump(preferencias, f)
    st.toast("Preferências salvas com sucesso!", icon="✅")

def converter_mes_romano(data_string):
    if pd.isna(data_string) or not str(data_string).strip(): return ""
    clean_str = str(data_string).replace(" 00:00:00", "").strip()
    try:
        data_obj = pd.to_datetime(clean_str, errors='coerce')
        if pd.isna(data_obj): return clean_str
        romanos = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]
        return f"{romanos[data_obj.month - 1]}.{data_obj.year}" 
    except:
        return clean_str

def converter_pdf_universal(input_docx, output_pdf):
    try:
        from docx2pdf import convert
        convert(input_docx, output_pdf)
        return True, "Microsoft Word"
    except Exception: pass 

    try:
        if platform.system() == 'Windows':
            lo_paths = [r"C:\Program Files\LibreOffice\program\soffice.exe", r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]
            soffice_path = next((p for p in lo_paths if os.path.exists(p)), None)
        elif platform.system() == 'Darwin': soffice_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        else: soffice_path = 'libreoffice'

        if soffice_path:
            subprocess.run([soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(os.path.abspath(output_pdf)) or '.', input_docx], check=True)
            return True, "LibreOffice"
    except Exception: pass 
    return False, None

# ==========================================
# 3. INTERFACE DE UPLOAD E MAPEAMENTO
# ==========================================
col1, col2 = st.columns([1, 2])
with col1:
    ficheiro_excel = st.file_uploader("1. Carregar Planilha Excel ou CSV", type=['xlsx', 'csv'])
    ficheiro_modelo = st.file_uploader("2. Carregar Modelo Word (.docx)", type=['docx'])

if ficheiro_excel and ficheiro_modelo:
    df = pd.read_csv(ficheiro_excel, dtype=str) if ficheiro_excel.name.endswith('.csv') else pd.read_excel(ficheiro_excel, dtype=str)
    colunas_excel = ["--- Não usar ---"] + df.columns.tolist()

    st.subheader("Mapeamento de Variáveis")
    prefs_salvas = carregar_preferencias()
    
    def obter_index(nome_campo):
        valor = prefs_salvas.get(nome_campo, "--- Não usar ---")
        return colunas_excel.index(valor) if valor in colunas_excel else 0

    col_a, col_b, col_c = st.columns(3)
    map_cols = {}
    with col_a:
        map_cols['Numero'] = st.selectbox("ID / Número", colunas_excel, index=obter_index('Numero'))
        map_cols['Duplicata'] = st.selectbox("Duplicata (A, B...)", colunas_excel, index=obter_index('Duplicata'))
        map_cols['Genero'] = st.selectbox("Gênero", colunas_excel, index=obter_index('Genero'))
    with col_b:
        map_cols['Especie'] = st.selectbox("Espécie", colunas_excel, index=obter_index('Especie'))
        map_cols['Autor'] = st.selectbox("Autor", colunas_excel, index=obter_index('Autor'))
        map_cols['Coletor'] = st.selectbox("Coletor", colunas_excel, index=obter_index('Coletor'))
    with col_c:
        map_cols['Determinador'] = st.selectbox("Determinador", colunas_excel, index=obter_index('Determinador'))
        map_cols['Data_Det'] = st.selectbox("Data de Determinação", colunas_excel, index=obter_index('Data_Det'))
        map_cols['Data_Original'] = st.selectbox("Data Original", colunas_excel, index=obter_index('Data_Original'))

    if st.button("💾 Salvar Preferências de Colunas"):
        guardar_preferencias(map_cols)

    # ==========================================
    # 4. PROCESSAMENTO E EXPORTAÇÃO
    # ==========================================
    st.divider()
    st.subheader("Opções de Exportação")
    
    formato_saida = st.radio("Escolha o formato do ficheiro final:", ["Word (.zip com etiquetas separadas)", "PDF (Arquivo Único)"], horizontal=True)

    if st.button(f"🚀 Gerar Etiquetas", type="primary"):
        st.info("A isolar as etiquetas e aplicar formatação cirúrgica...")
        
        itens_ignorados = []
        avisos = []
        etiquetas_validas = 0
        
        template_bytes = ficheiro_modelo.read()
        documentos_etiquetas = []

        if map_cols['Numero'] != "--- Não usar ---":
            agrupado = df.groupby(map_cols['Numero'])
            for numero, grupo_original in agrupado:
                
                # Inteligência de Agrupamento Blindada contra células vazias
                tem_duplicatas = False
                if map_cols['Duplicata'] != "--- Não usar ---":
                    dups_raw = grupo_original[map_cols['Duplicata']].tolist()
                    dups_validas = []
                    for d in dups_raw:
                        if pd.isna(d): continue
                        d_str = str(d).strip().lower()
                        # Filtro que ignora células que o Excel finge que estão preenchidas
                        if d_str in ['', 'nan', 'none', 'nat', '<na>', '-', '_', '.', '?', 'na', 'n/a', 'nd']: 
                            continue
                        dups_validas.append(d_str)
                    
                    if len(dups_validas) > 0:
                        tem_duplicatas = True
                        if len(dups_validas) != len(set(dups_validas)):
                            itens_ignorados.append(f"❌ O Número {numero} foi ignorado: Letras de duplicata repetidas na planilha ({', '.join(dups_validas)}).")
                            continue 
                        if len(dups_validas) > 3:
                            avisos.append(f"⚠️ Aviso no Número {numero}: Mais de 3 amostras na mesma etiqueta.")
                
                # Se tem letras de duplicata válidas, junta tudo. Se for vazio, separa cada linha numa etiqueta!
                subgrupos = [grupo_original] if tem_duplicatas else [grupo_original.iloc[[i]] for i in range(len(grupo_original))]

                for grupo in subgrupos:
                    etiquetas_validas += 1

                    def limpar_dado(coluna):
                        v = str(grupo.iloc[0][map_cols[coluna]]) if map_cols[coluna] != "--- Não usar ---" else ""
                        v = "" if v.lower() in ['nan', 'none', 'nat'] else v.strip()
                        if v.endswith(" 00:00:00"): v = v.replace(" 00:00:00", "")
                        return v

                    id_num = str(numero)
                    coletor = limpar_dado('Coletor')
                    det = limpar_dado('Determinador')
                    data_orig = limpar_dado('Data_Original')
                    data_romana = converter_mes_romano(data_orig)
                    
                    data_det_raw = limpar_dado('Data_Det')
                    ano_det = ""
                    if data_det_raw:
                        try: ano_det = str(pd.to_datetime(data_det_raw, dayfirst=True).year)
                        except: ano_det = data_det_raw[-4:]
                    
                    taxons_amostra = []
                    for _, row in grupo.iterrows():
                        gen = str(row[map_cols['Genero']]) if map_cols['Genero'] != "--- Não usar ---" else ""
                        esp = str(row[map_cols['Especie']]) if map_cols['Especie'] != "--- Não usar ---" else ""
                        aut = str(row[map_cols['Autor']]) if map_cols['Autor'] != "--- Não usar ---" else ""
                        dup = str(row[map_cols['Duplicata']]) if map_cols['Duplicata'] != "--- Não usar ---" else ""
                        
                        gen = "" if gen.lower() in ['nan', 'none'] else gen.strip()
                        esp = "" if esp.lower() in ['nan', 'none'] else esp.strip()
                        aut = "" if aut.lower() in ['nan', 'none'] else aut.strip()
                        
                        dup = "" if dup.lower() in ['nan', 'none', '<na>', '-', '_', '.', '?', 'na', 'n/a', 'nd'] else dup.strip().lower() 
                        
                        if esp.strip().lower() == "sp.": aut = ""
                        taxons_amostra.append({"genero": gen, "especie": esp, "autor": aut, "duplicata": dup})

                    # Clonagem Isolada para cada etiqueta
                    doc_temp = Document(io.BytesIO(template_bytes))
                    
                    for p_elm in doc_temp._element.xpath('.//w:p'):
                        p = Paragraph(p_elm, doc_temp)
                        if not p.text: continue
                        
                        def substituir_mantendo_estilo(paragrafo, chave, valor):
                            if chave not in paragrafo.text: return
                            for r in paragrafo.runs:
                                if chave in r.text:
                                    r.text = r.text.replace(chave, valor)
                                    return
                            estilo_base = paragrafo.runs[0]
                            nome_fonte = estilo_base.font.name
                            tamanho = estilo_base.font.size
                            negrito = estilo_base.bold
                            
                            novo_texto = paragrafo.text.replace(chave, valor)
                            paragrafo.clear()
                            novo_r = paragrafo.add_run(novo_texto)
                            if nome_fonte: novo_r.font.name = nome_fonte
                            if tamanho: novo_r.font.size = tamanho
                            if negrito: novo_r.bold = negrito

                        substituir_mantendo_estilo(p, '{{coletor}}', coletor)
                        substituir_mantendo_estilo(p, '{{determinador}}', det)
                        substituir_mantendo_estilo(p, '{{data_det}}', ano_det)
                        substituir_mantendo_estilo(p, '{{id}}', id_num)
                        substituir_mantendo_estilo(p, '{{data2}}', data_romana)
                        
                        if '{{conteudo}}' in p.text or '{{conteúdo}}' in p.text:
                            chave_conteudo = '{{conteúdo}}' if '{{conteúdo}}' in p.text else '{{conteudo}}'
                            
                            nome_fonte_conteudo = None
                            tamanho_fonte_conteudo = None
                            
                            for r in p.runs:
                                if chave_conteudo in r.text:
                                    nome_fonte_conteudo = r.font.name
                                    tamanho_fonte_conteudo = r.font.size
                                    r.text = r.text.replace(chave_conteudo, '') 
                                    break
                                    
                            for i, taxon in enumerate(taxons_amostra):
                                if taxon['genero']:
                                    r_gen = p.add_run(f"{taxon['genero']} ")
                                    if nome_fonte_conteudo: r_gen.font.name = nome_fonte_conteudo
                                    if tamanho_fonte_conteudo: r_gen.font.size = tamanho_fonte_conteudo
                                    r_gen.bold = True
                                    r_gen.italic = True
                                if taxon['especie']:
                                    r_esp = p.add_run(f"{taxon['especie']} ")
                                    if nome_fonte_conteudo: r_esp.font.name = nome_fonte_conteudo
                                    if tamanho_fonte_conteudo: r_esp.font.size = tamanho_fonte_conteudo
                                    r_esp.bold = True
                                    r_esp.italic = True
                                if taxon['autor']:
                                    r_aut = p.add_run(f"{taxon['autor']} ")
                                    if nome_fonte_conteudo: r_aut.font.name = nome_fonte_conteudo
                                    if tamanho_fonte_conteudo: r_aut.font.size = tamanho_fonte_conteudo
                                    r_aut.bold = True
                                    r_aut.italic = False
                                if taxon['duplicata']:
                                    r_dup = p.add_run(f"({taxon['duplicata']})")
                                    if nome_fonte_conteudo: r_dup.font.name = nome_fonte_conteudo
                                    if tamanho_fonte_conteudo: r_dup.font.size = tamanho_fonte_conteudo
                                    r_dup.bold = True
                                    r_dup.italic = False
                                if i < len(taxons_amostra) - 1:
                                    p.add_run("\n")
                                    
                    documentos_etiquetas.append((id_num, doc_temp))

        # Relatório Final
        if itens_ignorados:
            for ignorado in itens_ignorados: st.warning(ignorado)
        if avisos:
            for aviso in avisos: st.info(aviso)
            
        if etiquetas_validas > 0:
            if "Word" in formato_saida:
                st.success(f"Sucesso! {etiquetas_validas} etiquetas separadas estão a ser empacotadas num ficheiro ZIP. 📝")
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                    for i, (id_lbl, doc) in enumerate(documentos_etiquetas):
                        id_seguro = "".join([c for c in str(id_lbl) if c.isalnum() or c in " _-"])
                        if not id_seguro: id_seguro = f"etiqueta_{i+1}"
                        
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        zip_file.writestr(f"etiqueta_ID_{id_seguro}_{i+1}.docx", doc_buffer.getvalue())
                
                zip_buffer.seek(0)
                st.download_button(label="📥 Baixar Etiquetas (Word em ZIP)", data=zip_buffer, file_name="etiquetas_liquens.zip", mime="application/zip")
                
            elif "PDF" in formato_saida:
                st.info(f"A converter {etiquetas_validas} etiquetas de forma isolada (Isto pode demorar alguns segundos)...")
                
                try:
                    from PyPDF2 import PdfMerger
                    merger = PdfMerger()
                    sucesso_geral = True
                    barra_progresso = st.progress(0)
                    
                    for i, (id_lbl, doc) in enumerate(documentos_etiquetas):
                        id_seguro = "".join([c for c in str(id_lbl) if c.isalnum() or c in " _-"])
                        if not id_seguro: id_seguro = f"etiqueta_{i+1}"
                        
                        arquivo_temp_docx = f"temp_{id_seguro}_{i+1}.docx"
                        arquivo_temp_pdf = f"temp_{id_seguro}_{i+1}.pdf"
                        doc.save(arquivo_temp_docx)
                        
                        sucesso, motor = converter_pdf_universal(arquivo_temp_docx, arquivo_temp_pdf)
                        if sucesso:
                            with open(arquivo_temp_pdf, "rb") as f_pdf:
                                merger.append(io.BytesIO(f_pdf.read()))
                        else:
                            sucesso_geral = False
                            break
                        
                        if os.path.exists(arquivo_temp_docx): os.remove(arquivo_temp_docx)
                        if os.path.exists(arquivo_temp_pdf): os.remove(arquivo_temp_pdf)
                        
                        barra_progresso.progress((i + 1) / len(documentos_etiquetas))
                        
                    if sucesso_geral:
                        pdf_buffer = io.BytesIO()
                        merger.write(pdf_buffer)
                        merger.close()
                        pdf_buffer.seek(0)
                        st.success("Sucesso! O PDF unificado com os layouts intactos está pronto. 📄")
                        st.download_button(label="📥 Baixar Etiquetas (PDF Único)", data=pdf_buffer, file_name="etiquetas_liquens_final.pdf", mime="application/pdf")
                    else:
                        st.error("⚠️ Não foi possível converter para PDF automaticamente no seu computador.")
                        
                except ImportError:
                    st.error("⚠️ Erro! Precisa de instalar a ferramenta de PDF. Pare a aplicação (Ctrl+C), digite 'pip install PyPDF2' e rode novamente.")
        else:
            st.error("Nenhuma etiqueta válida pôde ser gerada.")